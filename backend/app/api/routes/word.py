"""
API Routes cho xử lý upload và submit form từ file Word/Document
"""

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Query
from fastapi.requests import Request
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import text, inspect
from typing import List, Dict, Iterable
import os
import json
import io
from urllib.parse import quote
from datetime import datetime
from pydantic import BaseModel
from docx import Document

from app.db.session import get_db
from app.db.models import WordTemplate, WordSubmission, Entry, Field, User, Form
from app.services.file_parser import FileParserFactory, FileField
from app.core.logger import logger
from app.core.file_utils import extract_clean_filename
from app.core.auth import get_current_user
from app.services.autofill.orchestrator import AutofillOrchestrator
from app.services.autofill.memory_service import UserMemoryService
import re

router = APIRouter(prefix="/api/word", tags=["word"])
autofill_orchestrator = AutofillOrchestrator()
template_form_instance_map: dict[int, int] = {}

SUBMISSION_META_KEY = "__autofill_meta__"

PLACEHOLDER_RE = re.compile(r'(\.{3,}|_{3,}|…{2,}|-{3,}|─{3,}|‒{3,}|–{3,}|—{3,}|\u2026{2,})')
DATE_TEMPLATE_RE = re.compile(
    r'ngày\s*[.\-_/…_‒–—]{1,}\s*tháng\s*[.\-_/…_‒–—]{1,}\s*năm\s*[.\-_/…_‒–—]{1,}',
    re.IGNORECASE
)

class WordDocxExportRequest(BaseModel):
    submission_ids: list[int]
    file_name: str | None = None
    user_id: int | None = None


class WordFieldLabelUpdateRequest(BaseModel):
    field_name: str
    new_label: str


def _iter_doc_paragraphs(doc: Document) -> Iterable:
    """Yield paragraph objects from body and tables in display order."""
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _normalize_export_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _split_submission_payload(raw_payload: object) -> tuple[dict, dict]:
    if not isinstance(raw_payload, dict):
        return {}, {}

    meta = raw_payload.get(SUBMISSION_META_KEY)
    if not isinstance(meta, dict):
        meta = {}

    visible_data: dict = {}
    for key, value in raw_payload.items():
        key_text = str(key)
        if key_text.startswith("__"):
            continue
        visible_data[key_text] = value

    return visible_data, meta


def _load_submission_payload(raw_json: str | None) -> tuple[dict, dict]:
    if not raw_json:
        return {}, {}

    try:
        parsed = json.loads(raw_json)
    except Exception:
        return {}, {}

    return _split_submission_payload(parsed)


def _compose_submission_payload(visible_data: dict, meta: dict | None = None) -> dict:
    payload: dict = dict(visible_data or {})
    if isinstance(meta, dict) and meta:
        payload[SUBMISSION_META_KEY] = meta
    return payload


def _parse_date_parts(raw_value: str) -> tuple[str, str, str] | None:
    raw = (raw_value or "").strip()
    if not raw:
        return None

    candidates: list[datetime] = []
    for date_fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            candidates.append(datetime.strptime(raw, date_fmt))
            break
        except ValueError:
            continue

    if not candidates:
        try:
            candidates.append(datetime.fromisoformat(raw))
        except ValueError:
            return None

    dt = candidates[0]
    return (f"{dt.day:02d}", f"{dt.month:02d}", f"{dt.year:04d}")


def _is_signature_date_line(text: str) -> bool:
    normalized = re.sub(r'\s+', ' ', (text or '').strip()).lower()
    if not normalized or not DATE_TEMPLATE_RE.search(normalized):
        return False
    if ":" in normalized or "：" in normalized:
        return False
    return not any(keyword in normalized for keyword in ("ngày nộp", "ngày ký", "ngày lập"))


def _pick_submission_date_value(data_map: dict, ordered_field_names: list[str]) -> str:
    for field_name in ordered_field_names:
        lowered = field_name.lower()
        if "ngay" in lowered and ("nop" in lowered or "don" in lowered or "ky" in lowered or "lap" in lowered):
            value = _normalize_export_value(data_map.get(field_name))
            if value:
                return value

    for key, value in data_map.items():
        lowered = str(key).lower()
        if "ngay" in lowered and ("nop" in lowered or "don" in lowered or "ky" in lowered or "lap" in lowered):
            cleaned = _normalize_export_value(value)
            if cleaned:
                return cleaned
    return ""


def _fill_submission_into_template_docx(template_path: str, data_map: dict, template_fields: list[dict]) -> tuple[Document, int]:
    """Fill values into uploaded .docx template and return replacement count."""
    doc = Document(template_path)
    ordered_field_names = [str(field.get("name", "")).strip() for field in template_fields if field.get("name")]
    values_in_order = [_normalize_export_value(data_map.get(name, "")) for name in ordered_field_names]
    value_idx = 0
    replaced_count = 0
    date_value = _pick_submission_date_value(data_map, ordered_field_names)

    for paragraph in _iter_doc_paragraphs(doc):
        original_text = paragraph.text or ""
        if not original_text:
            continue

        new_text = original_text

        # Keep signature footer date template stable, only replace if user has explicit date value.
        if _is_signature_date_line(original_text):
            parts = _parse_date_parts(date_value)
            if parts:
                dd, mm, yyyy = parts
                new_text = DATE_TEMPLATE_RE.sub(f"ngày {dd} tháng {mm} năm {yyyy}", new_text, count=1)
            if new_text != original_text:
                paragraph.text = new_text
                replaced_count += 1
            continue

        if not PLACEHOLDER_RE.search(new_text):
            continue

        def _replace_placeholder(match: re.Match) -> str:
            nonlocal value_idx
            if value_idx >= len(values_in_order):
                return match.group(0)
            value = values_in_order[value_idx]
            value_idx += 1
            return value if value else match.group(0)

        replaced_text = PLACEHOLDER_RE.sub(_replace_placeholder, new_text)
        if replaced_text != original_text:
            paragraph.text = replaced_text
            replaced_count += 1

    return doc, replaced_count


# Thư mục lưu file upload
UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)


def _sanitize_field_label(text: str) -> str:
    """Normalize legacy label text into a clean display label."""
    value = (text or "").strip()
    # Convert technical separators to spaces, remove punctuation/symbols,
    # then normalize to "Title Case" for clean UI/Word export display.
    value = re.sub(r'[_\-]+', ' ', value)
    value = re.sub(r'[^\w\s]', ' ', value)
    value = re.sub(r'\s+', ' ', value).strip()
    if not value:
        return ""
    return " ".join(word[:1].upper() + word[1:].lower() for word in value.split())


def _normalize_custom_field_label(text: str) -> str:
    """Normalize user custom label while preserving intended casing."""
    value = (text or "").strip()
    value = re.sub(r"\s+", " ", value).strip()
    return value


def _build_field_label_snapshot(template_fields: list[dict]) -> dict[str, str]:
    snapshot: dict[str, str] = {}
    for field_data in template_fields:
        field_name = str(field_data.get("name") or "").strip()
        if not field_name:
            continue

        field_label = _normalize_custom_field_label(str(field_data.get("label") or ""))
        snapshot[field_name] = field_label or field_name
    return snapshot


def _normalize_label_snapshot(raw_snapshot: object) -> dict[str, str]:
    if not isinstance(raw_snapshot, dict):
        return {}

    normalized: dict[str, str] = {}
    for raw_name, raw_label in raw_snapshot.items():
        field_name = str(raw_name or "").strip()
        if not field_name:
            continue

        field_label = _normalize_custom_field_label(str(raw_label or ""))
        if not field_label:
            field_label = field_name
        normalized[field_name] = field_label
    return normalized


def _classify_field_label_state(
    current_labels: dict[str, str],
    submitted_labels: dict[str, str]
) -> tuple[str, bool, int]:
    if not submitted_labels:
        return "unknown", False, 0

    diff_count = 0
    for field_name, current_label in current_labels.items():
        submitted_label = submitted_labels.get(field_name)
        if not submitted_label:
            continue
        if submitted_label != current_label:
            diff_count += 1

    if diff_count > 0:
        return "before_rename", True, diff_count
    return "after_rename", False, 0


def _to_field_name(label: str) -> str:
    """Create a snake_case field name from any label string."""
    value = (label or "").lower().strip()
    value = re.sub(r'[^\w\s]', ' ', value)
    value = re.sub(r'\s+', '_', value).strip('_')
    return value or "field"


def parse_template_fields(fields_json_raw: str | None) -> list[dict]:
    """Parse template fields_json and normalize legacy formats.

    Supported legacy formats:
    - list[dict] (current format)
    - list[str] (older placeholder-only format)
    - dict (single field)
    """
    if not fields_json_raw:
        return []

    try:
        payload = json.loads(fields_json_raw)
    except Exception:
        return []

    if isinstance(payload, dict):
        payload = [payload]
    if not isinstance(payload, list):
        return []

    normalized: list[dict] = []
    for idx, item in enumerate(payload):
        if isinstance(item, dict):
            raw_name = str(item.get("name") or "").strip()
            raw_label = _normalize_custom_field_label(str(item.get("label") or ""))
            if raw_label:
                label = raw_label
            else:
                fallback_label = raw_name or f"Field {idx + 1}"
                label = _sanitize_field_label(fallback_label) or f"Field {idx + 1}"

            name = _to_field_name(raw_name or label)
            field_type = item.get("field_type") or "text"
            order = item.get("order", idx)
        elif isinstance(item, str):
            label = _sanitize_field_label(item) or f"Field {idx + 1}"
            name = _to_field_name(label)
            field_type = "text"
            order = idx
        else:
            continue

        normalized.append({
            "name": str(name),
            "label": str(label),
            "field_type": str(field_type),
            "order": int(order) if isinstance(order, (int, float)) else idx,
            "label_updated_at": str(item.get("label_updated_at") or "").strip() if isinstance(item, dict) else ""
        })

    return normalized


def get_or_create_word_form_id(db: Session, user_id: int) -> int:
    """Get user-specific word form_id (create once if missing).

    Uses raw SQL with dynamic columns for compatibility with legacy schemas
    where ORM model columns may not exist yet.
    """
    try:
        ensure_forms_schema_compatibility(db)
    except Exception as schema_error:
        # Continue with dynamic SQL fallback even if ALTER TABLE is not permitted.
        logger.warning(f"Schema compatibility patch failed, using fallback form resolver: {schema_error}")

    bind = db.get_bind()
    inspector = inspect(bind)
    columns = {col["name"] for col in inspector.get_columns("forms")}

    select_sql = "SELECT id FROM forms WHERE user_id = :user_id"
    if "form_type" in columns:
        select_sql += " AND form_type = 'word'"
    select_sql += " ORDER BY id ASC LIMIT 1"

    row = db.execute(text(select_sql), {"user_id": user_id}).first()
    if row:
        return int(row[0])

    insert_columns: list[str] = ["user_id", "form_name"]
    insert_values_sql: list[str] = [":user_id", ":form_name"]
    params: dict[str, object] = {
        "user_id": user_id,
        "form_name": "Word Smart Form",
    }

    if "description" in columns:
        insert_columns.append("description")
        insert_values_sql.append(":description")
        params["description"] = "System form for Word template field mapping"
    if "form_type" in columns:
        insert_columns.append("form_type")
        insert_values_sql.append(":form_type")
        params["form_type"] = "word"
    if "is_template" in columns:
        insert_columns.append("is_template")
        insert_values_sql.append(":is_template")
        params["is_template"] = 1
    if "created_at" in columns:
        insert_columns.append("created_at")
        insert_values_sql.append("CURRENT_TIMESTAMP")
    if "updated_at" in columns:
        insert_columns.append("updated_at")
        insert_values_sql.append("CURRENT_TIMESTAMP")

    insert_sql = (
        f"INSERT INTO forms ({', '.join(insert_columns)}) "
        f"VALUES ({', '.join(insert_values_sql)})"
    )

    result = db.execute(text(insert_sql), params)
    db.commit()

    if result.lastrowid:
        return int(result.lastrowid)

    row = db.execute(text(select_sql), {"user_id": user_id}).first()
    if not row:
        raise HTTPException(status_code=500, detail="Không thể tạo form Word cho người dùng")
    return int(row[0])


def ensure_forms_schema_compatibility(db: Session) -> None:
    """Best-effort runtime compatibility for legacy forms table schema."""
    bind = db.get_bind()
    inspector = inspect(bind)

    if not inspector.has_table("forms"):
        # In case schema bootstrap was incomplete in some environments.
        Form.__table__.create(bind=bind, checkfirst=True)
        return

    try:
        existing_columns = {col["name"] for col in inspector.get_columns("forms")}
    except Exception as schema_error:
        logger.error(f"Unable to inspect forms table schema: {schema_error}")
        raise

    required_columns: dict[str, str] = {
        "description": "TEXT NULL",
        "form_type": "VARCHAR(50) NOT NULL DEFAULT 'standard'",
        "is_template": "BOOLEAN NOT NULL DEFAULT 0",
        "updated_at": "DATETIME NULL",
    }

    missing_columns = [name for name in required_columns if name not in existing_columns]
    if not missing_columns:
        return

    try:
        for col_name in missing_columns:
            ddl = required_columns[col_name]
            db.execute(text(f"ALTER TABLE forms ADD COLUMN {col_name} {ddl}"))

        # Backfill timestamps for legacy rows when column is newly added.
        if "updated_at" in missing_columns and "created_at" in existing_columns:
            db.execute(text("UPDATE forms SET updated_at = created_at WHERE updated_at IS NULL"))

        # Backfill defaults for old rows where these columns were absent.
        if "form_type" in missing_columns:
            db.execute(text("UPDATE forms SET form_type = 'standard' WHERE form_type IS NULL OR form_type = ''"))
        if "is_template" in missing_columns:
            db.execute(text("UPDATE forms SET is_template = 0 WHERE is_template IS NULL"))

        db.commit()
        logger.info(f"Patched legacy forms schema, added columns: {', '.join(missing_columns)}")
    except Exception as alter_error:
        db.rollback()
        alter_message = str(alter_error).lower()
        # Another worker/request may have added columns concurrently.
        if "duplicate column" in alter_message or "already exists" in alter_message:
            inspector = inspect(bind)
            refreshed_columns = {col["name"] for col in inspector.get_columns("forms")}
            still_missing = [name for name in required_columns if name not in refreshed_columns]
            if not still_missing:
                return
        logger.error(f"Failed to patch forms schema: {alter_error}")
        raise


def ensure_template_fields_exist(db: Session, form_id: int, fields_json: list[dict]) -> dict[str, int]:
    """Ensure template fields exist and return mapping by field_name -> field_id.

    Uses dynamic SQL to stay compatible with legacy `fields` table schemas.
    """
    bind = db.get_bind()
    inspector = inspect(bind)

    if not inspector.has_table("fields"):
        Field.__table__.create(bind=bind, checkfirst=True)

    columns = {col["name"] for col in inspector.get_columns("fields")}

    select_sql = "SELECT id, field_name FROM fields WHERE form_id = :form_id"
    existing_rows = db.execute(text(select_sql), {"form_id": form_id}).all()
    field_by_name: dict[str, int] = {str(row[1]): int(row[0]) for row in existing_rows if row[1]}

    created_any = False
    for idx, field_data in enumerate(fields_json):
        field_name = (field_data.get("name") or "").strip()
        if not field_name or field_name in field_by_name:
            continue

        insert_columns: list[str] = ["form_id", "field_name"]
        insert_values_sql: list[str] = [":form_id", ":field_name"]
        params: dict[str, object] = {
            "form_id": form_id,
            "field_name": field_name,
        }

        if "field_type" in columns:
            insert_columns.append("field_type")
            insert_values_sql.append(":field_type")
            params["field_type"] = field_data.get("field_type", "text")
        if "display_order" in columns:
            insert_columns.append("display_order")
            insert_values_sql.append(":display_order")
            params["display_order"] = field_data.get("order", idx)
        if "is_required" in columns:
            insert_columns.append("is_required")
            insert_values_sql.append(":is_required")
            params["is_required"] = 0
        if "validation_rules" in columns:
            insert_columns.append("validation_rules")
            insert_values_sql.append(":validation_rules")
            params["validation_rules"] = None
        if "placeholder" in columns:
            insert_columns.append("placeholder")
            insert_values_sql.append(":placeholder")
            params["placeholder"] = field_data.get("label", field_name)
        if "created_at" in columns:
            insert_columns.append("created_at")
            insert_values_sql.append("CURRENT_TIMESTAMP")

        insert_sql = (
            f"INSERT INTO fields ({', '.join(insert_columns)}) "
            f"VALUES ({', '.join(insert_values_sql)})"
        )
        result = db.execute(text(insert_sql), params)
        created_any = True

        if result.lastrowid:
            field_by_name[field_name] = int(result.lastrowid)

    if created_any:
        db.commit()

    # Fill any unresolved IDs (e.g. driver did not return lastrowid)
    if any(name not in field_by_name for name in [f.get("name", "") for f in fields_json if f.get("name")]):
        refreshed_rows = db.execute(text(select_sql), {"form_id": form_id}).all()
        refreshed_map = {str(row[1]): int(row[0]) for row in refreshed_rows if row[1]}
        field_by_name.update(refreshed_map)

    return field_by_name


@router.post("/autofill/{template_id}")
async def run_word_template_autofill(
    template_id: int,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Run LLM autofill for a Word template canonical form instance."""
    effective_user_id = resolve_effective_user_id(current_user, user_id)
    template = db.query(WordTemplate).filter(
        WordTemplate.id == template_id,
        WordTemplate.user_id == effective_user_id,
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template không tồn tại")

    form_instance_id = template_form_instance_map.get(template_id)
    if not form_instance_id:
        if not os.path.exists(template.file_path):
            raise HTTPException(status_code=404, detail="File template không tồn tại")
        _, form_instance_id = await autofill_orchestrator.parse_and_prepare_schema(
            db=db,
            user_id=effective_user_id,
            file_path=template.file_path,
            source_type="word",
            source_ref=str(template.id),
            original_filename=template.original_filename,
        )
        template_form_instance_map[template_id] = form_instance_id

    result = await autofill_orchestrator.run_autofill(
        db=db,
        user_id=effective_user_id,
        form_instance_id=form_instance_id,
    )
    return {
        "status": "success",
        "template_id": template_id,
        "form_instance_id": form_instance_id,
        **result,
    }


@router.get("/autofill-status/{template_id}")
async def get_word_autofill_status(
    template_id: int,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Quick endpoint to help UI know whether LLM form instance is ready."""
    effective_user_id = resolve_effective_user_id(current_user, user_id)
    template = db.query(WordTemplate).filter(
        WordTemplate.id == template_id,
        WordTemplate.user_id == effective_user_id,
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template không tồn tại")
    return {
        "status": "success",
        "template_id": template_id,
        "form_instance_id": template_form_instance_map.get(template_id),
        "ready": bool(template_form_instance_map.get(template_id)),
    }


def _sync_word_field_placeholder(
    db: Session,
    form_id: int,
    field_name: str,
    field_label: str
) -> None:
    """Best effort: keep `fields.placeholder` aligned with label edits."""
    bind = db.get_bind()
    inspector = inspect(bind)
    if not inspector.has_table("fields"):
        return

    columns = {col["name"] for col in inspector.get_columns("fields")}
    assignments: list[str] = []
    params: dict[str, object] = {
        "form_id": form_id,
        "field_name": field_name,
    }

    if "placeholder" in columns:
        assignments.append("placeholder = :placeholder")
        params["placeholder"] = field_label
    if "updated_at" in columns:
        assignments.append("updated_at = CURRENT_TIMESTAMP")

    if not assignments:
        return

    db.execute(
        text(
            f"UPDATE fields SET {', '.join(assignments)} "
            "WHERE form_id = :form_id AND field_name = :field_name"
        ),
        params
    )
    db.commit()


def resolve_effective_user_id(current_user: User, requested_user_id: int | None) -> int:
    """Resolve target user scope. Non-admin users can only access their own data."""
    if requested_user_id is None:
        return current_user.id

    if requested_user_id <= 0:
        raise HTTPException(status_code=400, detail="user_id phải là số dương")

    if not current_user.is_admin and requested_user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Bạn không có quyền truy cập dữ liệu của người dùng khác")

    return requested_user_id


@router.post("/upload")
async def upload_word_template(
    file: UploadFile = File(...),
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Upload file (Word, PDF, Excel, CSV, Text) và parse thành template
    
    Hỗ trợ các định dạng:
    - .docx (Word Document)
    - .pdf (PDF File)
    - .xlsx, .xls (Excel File)
    - .csv (CSV File)
    - .txt (Text File)
    """
    
    # Kiểm tra file extension
    file_ext = os.path.splitext(file.filename)[1].lower()
    if not FileParserFactory.is_supported(file.filename):
        supported = ', '.join(FileParserFactory.get_supported_extensions())
        raise HTTPException(
            status_code=400, 
            detail=f"Không hỗ trợ định dạng file: {file_ext}. Các định dạng được hỗ trợ: {supported}"
        )
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)

        # Lưu file
        file_path = os.path.join(UPLOAD_DIR, f"{effective_user_id}_{datetime.now().timestamp()}_{file.filename}")
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Parse file bằng parser phù hợp
        parser = FileParserFactory.create_parser(file_path)
        fields = parser.parse()
        metadata = parser.get_metadata()
        
        # Nếu vẫn không tìm được field nào, tạo một default field từ tên file
        if not fields:
            # Tạo field tạm thời dựa trên tên file
            default_field_name = os.path.splitext(file.filename)[0]
            fields = [FileField(
                name=default_field_name.lower().replace(' ', '_'),
                field_type="text",
                label=f"Nội dung từ {file.filename}",
                order=0
            )]
        
        # Lưu template vào database
        template = WordTemplate(
            user_id=effective_user_id,
            template_name=metadata.get("title", file.filename),
            file_path=file_path,
            original_filename=file.filename,
            fields_json=json.dumps([f.to_dict() for f in fields])
        )
        db.add(template)
        db.commit()
        db.refresh(template)

        fields_json = parse_template_fields(template.fields_json)
        word_form_id = get_or_create_word_form_id(db, effective_user_id)
        ensure_template_fields_exist(db, word_form_id, fields_json)
        form_instance_id = None
        try:
            _, form_instance_id = await autofill_orchestrator.parse_and_prepare_schema(
                db=db,
                user_id=effective_user_id,
                file_path=file_path,
                source_type="word",
                source_ref=str(template.id),
                original_filename=file.filename,
            )
            if form_instance_id:
                template_form_instance_map[template.id] = form_instance_id
        except Exception as parse_error:
            logger.warning(f"Unable to build canonical form instance for Word template {template.id}: {parse_error}")
        
        auto_generated = len(fields) == 1 and "nội dung" in [f.to_dict() for f in fields][0].get("label", "").lower()
        
        return {
            "status": "success",
            "template_id": template.id,
            "template_name": template.template_name,
            "file_type": file_ext,
            "fields_count": len(fields),
            "fields": [f.to_dict() for f in fields],
            "form_instance_id": form_instance_id,
            "auto_generated_fields": auto_generated,
            "message": "Upload và parse thành công" if not auto_generated else "Upload thành công. Không tìm thấy trường cấu trúc, đã tạo trường mặc định."
        }
    
    except Exception as e:
        logger.error(f"Lỗi upload file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/supported-formats")
async def get_supported_formats():
    """Lấy danh sách các định dạng file được hỗ trợ"""
    return {
        "supported_extensions": FileParserFactory.get_supported_extensions(),
        "description": {
            ".docx": "Word Document",
            ".pdf": "PDF File", 
            ".xlsx": "Excel Spreadsheet (2007+)",
            ".xls": "Excel Spreadsheet (97-2003)",
            ".csv": "Comma-Separated Values",
            ".txt": "Plain Text File"
        }
    }


@router.get("/templates")
async def get_user_templates(
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Lấy danh sách template của user"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    templates = db.query(WordTemplate).filter(WordTemplate.user_id == effective_user_id).all()
    
    return {
        "templates": [
            {
                "id": t.id,
                "name": t.template_name,
                "filename": extract_clean_filename(t.original_filename),
                "fields_count": len(parse_template_fields(t.fields_json)),
                "created_at": t.created_at.isoformat(),
                "submissions_count": len(t.submissions)
            }
            for t in templates
        ]
    }


@router.get("/template/{template_id}")
async def get_template_detail(
    template_id: int,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Lấy chi tiết template"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    template = db.query(WordTemplate).filter(
        WordTemplate.id == template_id,
        WordTemplate.user_id == effective_user_id
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template không tồn tại")

    if not os.path.exists(template.file_path):
        raise HTTPException(status_code=404, detail="File template không tồn tại trên server")

    form_id = get_or_create_word_form_id(db, effective_user_id)
    
    fields_json = parse_template_fields(template.fields_json)
    field_by_name = ensure_template_fields_exist(db, form_id, fields_json)

    # Enrich fields với database field IDs
    enriched_fields = []
    for idx, field_data in enumerate(fields_json):
        field_name = field_data.get("name", "")
        field_label = field_data.get("label", field_name)
        field_id = field_by_name.get(field_name, -1)
        
        enriched_fields.append({
            **field_data,
            "field_id": field_id,
            "field_index": idx,
            "field_label": field_label
        })
    
    try:
        submissions_count = len(template.submissions)
    except Exception as e:
        logger.error(f"Error getting submissions count: {str(e)}")
        submissions_count = 0
    
    return {
        "id": template.id,
        "name": template.template_name,
        "filename": extract_clean_filename(template.original_filename),
        "fields": enriched_fields,
        "form_id": form_id,
        "form_instance_id": template_form_instance_map.get(template.id),
        "created_at": template.created_at.isoformat() if template.created_at else None,
        "submissions_count": submissions_count
    }


@router.put("/template/{template_id}/field-label")
async def update_template_field_label(
    template_id: int,
    request: WordFieldLabelUpdateRequest,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Update display label for a template field without changing submitted keys."""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    template = db.query(WordTemplate).filter(
        WordTemplate.id == template_id,
        WordTemplate.user_id == effective_user_id
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template không tồn tại")

    target_field_name = str(request.field_name or "").strip()
    if not target_field_name:
        raise HTTPException(status_code=400, detail="field_name không hợp lệ")

    new_label = _normalize_custom_field_label(request.new_label)
    if not new_label:
        raise HTTPException(status_code=400, detail="new_label không hợp lệ")

    fields_json = parse_template_fields(template.fields_json)
    if not fields_json:
        raise HTTPException(status_code=404, detail="Template không có trường để cập nhật")

    target_found = False
    for field_data in fields_json:
        field_name = str(field_data.get("name") or "").strip()
        if field_name != target_field_name:
            continue

        field_data["label"] = new_label
        field_data["label_updated_at"] = datetime.utcnow().isoformat()
        target_found = True
        break

    if not target_found:
        raise HTTPException(status_code=404, detail="Không tìm thấy trường cần đổi tên")

    template.fields_json = json.dumps(fields_json, ensure_ascii=False)
    db.commit()
    db.refresh(template)

    try:
        form_id = get_or_create_word_form_id(db, effective_user_id)
        _sync_word_field_placeholder(db, form_id, target_field_name, new_label)
    except Exception as sync_error:
        db.rollback()
        logger.warning(f"Unable to sync word field placeholder for '{target_field_name}': {sync_error}")
        form_id = get_or_create_word_form_id(db, effective_user_id)

    enriched_fields = parse_template_fields(template.fields_json)
    field_by_name = ensure_template_fields_exist(db, form_id, enriched_fields)
    response_fields = []
    for idx, field_data in enumerate(enriched_fields):
        field_name = str(field_data.get("name") or "").strip()
        field_label = str(field_data.get("label") or field_name).strip()
        response_fields.append({
            **field_data,
            "field_id": field_by_name.get(field_name, -1),
            "field_index": idx,
            "field_label": field_label,
        })

    return {
        "status": "success",
        "template_id": template.id,
        "field_name": target_field_name,
        "field_label": new_label,
        "fields": response_fields,
        "message": "Đã cập nhật tên trường thành công"
    }


@router.post("/submit")
async def submit_form(
    request: Request,
    template_id: int = Query(...),
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Submit form data cho template"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)
    form_id = get_or_create_word_form_id(db, effective_user_id)
    logger.info(f"submit_form called: template_id={template_id}, user_id={effective_user_id}, form_id={form_id}")
    
    template = db.query(WordTemplate).filter(
        WordTemplate.id == template_id,
        WordTemplate.user_id == effective_user_id
    ).first()
    if not template:
        logger.error(f"Template {template_id} not found")
        raise HTTPException(status_code=404, detail="Template không tồn tại")
    
    # Get JSON data from request body
    try:
        request_data = await request.json()
        if not isinstance(request_data, dict):
            request_data = {}
    except Exception as e:
        logger.error(f"Error parsing JSON: {str(e)}", exc_info=True)
        request_data = {}

    data, submission_meta = _split_submission_payload(request_data)

    template_fields_for_snapshot = parse_template_fields(template.fields_json)
    label_snapshot = _build_field_label_snapshot(template_fields_for_snapshot)
    if label_snapshot:
        if not isinstance(submission_meta, dict):
            submission_meta = {}
        submission_meta["field_label_snapshot"] = label_snapshot
    logger.info(f"Received form data: {data}")
    
    try:
        # Lưu submission
        payload_to_store = _compose_submission_payload(data, submission_meta)
        submission = WordSubmission(
            template_id=template_id,
            user_id=effective_user_id,
            submission_data=json.dumps(payload_to_store or {}, ensure_ascii=False)
        )
        db.add(submission)
        db.commit()
        db.refresh(submission)
        
        logger.info(f"Submission saved: {submission.id}")
        
        # Lưu entries
        try:
            fields_json = template_fields_for_snapshot
            logger.info(f"Template has {len(fields_json)} fields")
            field_by_name = ensure_template_fields_exist(db, form_id, fields_json)
            entries_to_insert: list[Entry] = []
            
            for field_data in fields_json:
                field_name = field_data.get("name", "")
                logger.info(f"Processing field: {field_name}")
                
                # Lấy giá trị từ form submit
                value = data.get(field_name, "").strip()
                if not value:
                    logger.info(f"  Field {field_name} has no value, skipping")
                    continue
                
                logger.info(f"  Field {field_name} value: '{value}'")
                
                field_id = field_by_name.get(field_name)
                if not field_id:
                    logger.warning(f"  Field mapping missing for {field_name}, skipping")
                    continue
                
                # Lưu entry
                entry = Entry(
                    user_id=effective_user_id,
                    field_id=field_id,
                    form_id=form_id,
                    value=value
                )
                entries_to_insert.append(entry)

            if entries_to_insert:
                db.add_all(entries_to_insert)
                db.commit()
                logger.info(f"Saved {len(entries_to_insert)} entries for submission {submission.id}")
                for entry in entries_to_insert:
                    try:
                        normalized_field = db.query(Field).filter(Field.id == entry.field_id).first()
                        field_key = normalized_field.field_name if normalized_field else ""
                        if field_key:
                            UserMemoryService.upsert_memory_value(
                                db=db,
                                user_id=effective_user_id,
                                field_key=field_key,
                                value_text=entry.value,
                                memory_type="entry",
                                source_ref=f"word_submission:{submission.id}",
                                confidence=0.7,
                                is_confirmed=False,
                            )
                    except Exception as memory_error:
                        logger.warning(f"Unable to sync Word entry to memory: {memory_error}")
                db.commit()
        
        except Exception as e:
            logger.error(f"Error saving entries: {str(e)}", exc_info=True)
            # Không throw error, submission đã lưu
        
        logger.info(f"Submit completed for submission {submission.id}")
        return {
            "status": "success",
            "submission_id": submission.id,
            "history_group_id": str(submission_meta.get("group_id") or "").strip() or None,
            "message": "Submit thành công"
        }
    
    except Exception as e:
        logger.error(f"Error in submit_form: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/submissions")
async def get_submissions(
    template_id: int = Query(None),
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Lấy danh sách submission"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    query = db.query(WordSubmission).filter(WordSubmission.user_id == effective_user_id)
    
    if template_id:
        query = query.filter(WordSubmission.template_id == template_id)
    
    submissions = query.all()
    
    normalized_items = []
    for s in submissions:
        data_map, meta = _load_submission_payload(s.submission_data)
        template_fields = parse_template_fields(s.template.fields_json if s.template else None)
        current_label_snapshot = _build_field_label_snapshot(template_fields)
        submitted_label_snapshot = _normalize_label_snapshot(meta.get("field_label_snapshot") if isinstance(meta, dict) else None)
        label_state, has_renamed_fields, renamed_field_count = _classify_field_label_state(
            current_label_snapshot,
            submitted_label_snapshot
        )
        normalized_items.append(
            {
                "id": s.id,
                "template_id": s.template_id,
                "template_name": s.template.template_name,
                "data": data_map,
                "field_label_state": label_state,
                "has_renamed_fields": has_renamed_fields,
                "renamed_field_count": renamed_field_count,
                "history_group_id": str(meta.get("group_id") or "").strip() or None,
                "created_at": s.created_at.isoformat()
            }
        )

    return {"submissions": normalized_items}


@router.get("/submission/{submission_id}")
async def get_submission_detail(
    submission_id: int,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Lấy chi tiết một submission"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    submission = db.query(WordSubmission).filter(
        WordSubmission.id == submission_id,
        WordSubmission.user_id == effective_user_id
    ).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission không tồn tại")
    
    data_map, meta = _load_submission_payload(submission.submission_data)
    template_fields = parse_template_fields(submission.template.fields_json if submission.template else None)
    template_field_map: dict[str, dict] = {}
    current_label_snapshot = _build_field_label_snapshot(template_fields)
    submitted_label_snapshot = _normalize_label_snapshot(meta.get("field_label_snapshot") if isinstance(meta, dict) else None)
    label_state, has_renamed_fields, renamed_field_count = _classify_field_label_state(
        current_label_snapshot,
        submitted_label_snapshot
    )
    if template_fields:
        form_id = get_or_create_word_form_id(db, effective_user_id)
        field_by_name = ensure_template_fields_exist(db, form_id, template_fields)
        for field_data in template_fields:
            field_name = str(field_data.get("name") or "").strip()
            if not field_name:
                continue
            template_field_map[field_name] = {
                "field_id": field_by_name.get(field_name, -1),
                "field_label": str(field_data.get("label") or field_name).strip(),
                "field_type": str(field_data.get("field_type") or "text").strip() or "text",
            }

    return {
        "id": submission.id,
        "template_id": submission.template_id,
        "template_name": submission.template.template_name,
        "data": data_map,
        "template_fields": template_field_map,
        "field_label_state": label_state,
        "has_renamed_fields": has_renamed_fields,
        "renamed_field_count": renamed_field_count,
        "submitted_field_labels": submitted_label_snapshot,
        "current_field_labels": current_label_snapshot,
        "history_group_id": str(meta.get("group_id") or "").strip() or None,
        "created_at": submission.created_at.isoformat(),
        "updated_at": submission.created_at.isoformat()
    }


@router.post("/export-docx")
@router.post("/export_docx")
async def export_submissions_docx(
    payload: WordDocxExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Export .docx, ưu tiên giữ nguyên layout template gốc khi có thể."""
    if not payload.submission_ids:
        raise HTTPException(status_code=400, detail="Vui lòng chọn ít nhất một submission để xuất")

    effective_user_id = resolve_effective_user_id(current_user, payload.user_id)
    normalized_ids = sorted({sid for sid in payload.submission_ids if isinstance(sid, int) and sid > 0})
    if not normalized_ids:
        raise HTTPException(status_code=400, detail="Danh sách submission không hợp lệ")

    submissions = db.query(WordSubmission).filter(
        WordSubmission.user_id == effective_user_id,
        WordSubmission.id.in_(normalized_ids)
    ).all()

    if not submissions:
        raise HTTPException(status_code=404, detail="Không tìm thấy submission phù hợp để xuất")

    submissions_by_id = {int(item.id): item for item in submissions}
    ordered_submissions = [submissions_by_id[sid] for sid in normalized_ids if sid in submissions_by_id]

    # Best-effort: nếu export 1 submission từ template .docx gốc thì điền trực tiếp vào template,
    # nhờ đó giữ nguyên định dạng giống file người dùng đã upload.
    if len(ordered_submissions) == 1:
        submission = ordered_submissions[0]
        template = submission.template
        template_path = (template.file_path or "") if template else ""
        if template and template_path.lower().endswith(".docx") and os.path.exists(template_path):
            data_map, _ = _load_submission_payload(submission.submission_data)
            template_fields = parse_template_fields(template.fields_json if template else None)

            try:
                rendered_doc, replaced_count = _fill_submission_into_template_docx(
                    template_path=template_path,
                    data_map=data_map,
                    template_fields=template_fields
                )
                if replaced_count <= 0:
                    raise ValueError("Template has no usable placeholders, fallback to auto-formatted export")
                buffer = io.BytesIO()
                rendered_doc.save(buffer)
                buffer.seek(0)

                requested_name = (payload.file_name or template.original_filename or "word_form_export").strip()
                safe_name = re.sub(r'[\\/:*?"<>|]+', "_", requested_name) or "word_form_export"
                if not safe_name.lower().endswith(".docx"):
                    safe_name = f"{safe_name}.docx"

                encoded_name = quote(safe_name)
                headers = {
                    "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"
                }
                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers=headers
                )
            except Exception as render_error:
                logger.warning(f"Template-preserving export failed, fallback to plain export: {render_error}")

    doc = Document()

    for idx, submission in enumerate(ordered_submissions):
        data_map, _ = _load_submission_payload(submission.submission_data)
        template_fields = parse_template_fields(submission.template.fields_json if submission.template else None)
        ordered_field_names = [str(field.get("name", "")).strip() for field in template_fields if field.get("name")]
        label_by_name = {
            str(field.get("name", "")).strip(): _sanitize_field_label(str(field.get("label", "")).strip())
            for field in template_fields
            if field.get("name")
        }

        emitted_fields: set[str] = set()
        for field_name in ordered_field_names:
            if field_name not in data_map:
                continue
            value = str(data_map.get(field_name, "")).strip()
            if value == "":
                continue

            clean_label = label_by_name.get(field_name) or _sanitize_field_label(field_name) or field_name
            doc.add_paragraph(f"{clean_label}: {value}")
            emitted_fields.add(field_name)

        for extra_key, extra_value in data_map.items():
            field_name = str(extra_key).strip()
            if not field_name or field_name in emitted_fields:
                continue
            value = str(extra_value or "").strip()
            if value == "":
                continue

            clean_label = _sanitize_field_label(field_name) or field_name
            doc.add_paragraph(f"{clean_label}: {value}")

        if idx < len(ordered_submissions) - 1:
            doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    requested_name = (payload.file_name or "word_form_export").strip()
    safe_name = re.sub(r'[\\/:*?"<>|]+', "_", requested_name) or "word_form_export"
    if not safe_name.lower().endswith(".docx"):
        safe_name = f"{safe_name}.docx"

    encoded_name = quote(safe_name)
    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"
    }
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@router.put("/submission/{submission_id}")
async def update_submission(
    submission_id: int,
    request: Request,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Cập nhật dữ liệu submission"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    submission = db.query(WordSubmission).filter(
        WordSubmission.id == submission_id,
        WordSubmission.user_id == effective_user_id
    ).first()
    
    if not submission:
        raise HTTPException(status_code=404, detail="Submission không tồn tại")
    
    try:
        # Đọc JSON body từ request
        incoming_payload = await request.json()
        if incoming_payload is None or not isinstance(incoming_payload, dict):
            incoming_payload = {}

        _, current_meta = _load_submission_payload(submission.submission_data)
        incoming_data_map, incoming_meta = _split_submission_payload(incoming_payload)
        effective_meta = dict(current_meta or {})
        if incoming_meta:
            effective_meta.update(incoming_meta)

        # Khi người dùng chỉnh sửa submission, cập nhật mốc label snapshot theo template hiện tại
        # để phân loại đúng "sau đổi tên" cho các trang đã được chỉnh sửa.
        template_fields_for_snapshot = parse_template_fields(submission.template.fields_json if submission.template else None)
        label_snapshot = _build_field_label_snapshot(template_fields_for_snapshot)
        if label_snapshot:
            effective_meta["field_label_snapshot"] = label_snapshot

        submission.submission_data = json.dumps(
            _compose_submission_payload(incoming_data_map, effective_meta),
            ensure_ascii=False
        )
        db.commit()
        db.refresh(submission)
        
        return {
            "status": "success",
            "message": "Cập nhật submission thành công",
            "id": submission.id,
            "data": incoming_data_map,
            "history_group_id": str(effective_meta.get("group_id") or "").strip() or None
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Lỗi cập nhật: {str(e)}")


@router.delete("/submission/{submission_id}")
async def delete_submission(
    submission_id: int,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Xóa submission"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    submission = db.query(WordSubmission).filter(
        WordSubmission.id == submission_id,
        WordSubmission.user_id == effective_user_id
    ).first()
    
    if not submission:
        raise HTTPException(status_code=404, detail="Submission không tồn tại")
    
    try:
        db.delete(submission)
        db.commit()
        
        return {
            "status": "success",
            "message": "Xóa submission thành công"
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Lỗi xóa: {str(e)}")


@router.delete("/template/{template_id}")
async def delete_template(
    template_id: int,
    user_id: int | None = Query(None, description="ID của user (optional, admin only)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Xóa template"""
    effective_user_id = resolve_effective_user_id(current_user, user_id)

    template = db.query(WordTemplate).filter(
        WordTemplate.id == template_id,
        WordTemplate.user_id == effective_user_id
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template không tồn tại")
    
    # Xóa file
    if os.path.exists(template.file_path):
        os.remove(template.file_path)
    
    # Xóa record
    db.delete(template)
    db.commit()
    
    return {"status": "success", "message": "Xóa template thành công"}
