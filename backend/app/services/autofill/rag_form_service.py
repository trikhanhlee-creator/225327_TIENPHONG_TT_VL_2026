from __future__ import annotations

import os
import re
from typing import Any

from docx import Document
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.logger import logger
from app.services.autofill.contracts import CanonicalFormField, MemoryCandidate
from app.services.autofill.memory_chunk_indexer import StandaloneTextIndexer
from app.services.autofill.memory_retrieval_agent import LLMMemoryRetrievalAgent
from app.services.suggestion_value_filters import is_valid_field_suggestion_value


def _normalize_field_key(raw: str) -> str:
    return str(raw or "").strip().lower().replace(" ", "_")


def extract_indexable_text(file_path: str) -> str:
    """Extract plain text from uploaded files for RAG indexing (best-effort)."""
    ext = os.path.splitext(file_path or "")[1].lower()
    if ext not in (".docx", ".doc") or not os.path.exists(file_path):
        return ""

    docx_path = file_path
    if ext == ".doc":
        try:
            from app.services.doc_converter import ensure_docx_for_processing

            docx_path = ensure_docx_for_processing(file_path)
        except Exception as exc:
            logger.warning(f"Unable to convert .doc for RAG text extraction: {exc}")
            return ""

    try:
        doc = Document(docx_path)
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = (paragraph.text or "").strip()
                        if text:
                            parts.append(text)
        return "\n".join(parts).strip()
    except Exception as exc:
        logger.warning(f"Unable to extract docx text for RAG: {exc}")
        return ""


class RagFormService:
    """RAG helpers for parse, UI hints, and Word export fill."""

    def __init__(self, retrieval_agent: LLMMemoryRetrievalAgent | None = None) -> None:
        self._retrieval = retrieval_agent or LLMMemoryRetrievalAgent()

    def index_uploaded_file(
        self,
        db: Session,
        *,
        user_id: int,
        file_path: str,
        source_ref: str,
    ) -> int:
        if not settings.RAG_ENABLED:
            return 0
        text = extract_indexable_text(file_path)
        if len(text) < 40:
            return 0
        try:
            count = StandaloneTextIndexer.index_plaintext(
                db,
                user_id=user_id,
                text=text,
                source_ref=f"upload:{source_ref}",
            )
            if count:
                db.commit()
                logger.info(f"RAG indexed {count} chunks for upload source_ref={source_ref}")
            return count
        except Exception as exc:
            logger.warning(f"RAG upload indexing failed: {exc}")
            db.rollback()
            return 0

    def build_field_hints(
        self,
        db: Session,
        *,
        user_id: int,
        fields: list[CanonicalFormField],
        top_k: int = 3,
        min_confidence: float = 0.45,
    ) -> dict[str, dict[str, Any]]:
        if not settings.RAG_ENABLED or not fields:
            return {}

        hints: dict[str, dict[str, Any]] = {}
        for field in fields:
            field_key = _normalize_field_key(field.field_key)
            if not field_key:
                continue
            candidates = self._retrieval.retrieve_for_field(
                db=db,
                user_id=user_id,
                field=field,
                top_k=max(1, top_k),
            )
            payload = self._candidate_to_hint(field_key, candidates, min_confidence=min_confidence)
            if payload:
                hints[field_key] = payload
        return hints

    def suggest_values_for_field_name(
        self,
        db: Session,
        *,
        user_id: int,
        field_name: str,
        field_label: str | None = None,
        field_type: str = "text",
        top_k: int = 5,
        min_confidence: float = 0.4,
    ) -> list[dict[str, Any]]:
        if not settings.RAG_ENABLED:
            return []

        field_key = _normalize_field_key(field_name)
        if not field_key:
            return []

        field = CanonicalFormField(
            field_key=field_key,
            label=(field_label or field_name or field_key).strip(),
            field_type=field_type or "text",
        )
        candidates = self._retrieval.retrieve_for_field(
            db=db,
            user_id=user_id,
            field=field,
            top_k=max(1, min(top_k, 10)),
        )
        out: list[dict[str, Any]] = []
        seen: set[str] = set()
        for candidate in candidates:
            value = (candidate.value or "").strip()
            if not value or value in seen:
                continue
            if not is_valid_field_suggestion_value(value):
                continue
            if candidate.confidence < min_confidence:
                continue
            seen.add(value)
            tier = int(candidate.metadata.get("tier", 99))
            out.append(
                {
                    "value": value,
                    "confidence": float(candidate.confidence),
                    "source": str(candidate.memory_type or "rag"),
                    "tier": tier,
                    "score": float(candidate.metadata.get("composite_score", candidate.score)),
                }
            )
        return out

    def enrich_data_map(
        self,
        db: Session,
        *,
        user_id: int,
        template_fields: list[dict],
        data_map: dict,
        only_empty: bool = True,
        min_confidence: float = 0.5,
    ) -> tuple[dict, int]:
        if not settings.RAG_ENABLED or not template_fields:
            return dict(data_map or {}), 0

        enriched = dict(data_map or {})
        filled = 0
        for field_data in template_fields:
            name = str(field_data.get("name") or "").strip()
            if not name:
                continue
            current = str(enriched.get(name) or "").strip()
            if only_empty and current:
                continue

            field = CanonicalFormField(
                field_key=_normalize_field_key(name),
                label=str(field_data.get("label") or name).strip(),
                field_type=str(field_data.get("field_type") or "text").strip() or "text",
            )
            candidates = self._retrieval.retrieve_for_field(
                db=db,
                user_id=user_id,
                field=field,
                top_k=3,
            )
            hint = self._candidate_to_hint(field.field_key, candidates, min_confidence=min_confidence)
            if not hint:
                continue
            enriched[name] = hint["value"]
            filled += 1
        return enriched, filled

    @staticmethod
    def merge_hints_into_template_fields(
        fields: list[dict],
        hints: dict[str, dict[str, Any]],
    ) -> list[dict]:
        if not hints:
            return fields

        merged: list[dict] = []
        for item in fields:
            row = dict(item)
            key = _normalize_field_key(str(row.get("name") or row.get("field_key") or ""))
            hint = hints.get(key)
            if hint:
                row["suggested_value"] = hint.get("value", "")
                row["suggestion_confidence"] = hint.get("confidence", 0.0)
                row["suggestion_source"] = hint.get("source", "rag")
            merged.append(row)
        return merged

    def _candidate_to_hint(
        self,
        field_key: str,
        candidates: list[MemoryCandidate],
        *,
        min_confidence: float,
    ) -> dict[str, Any] | None:
        for candidate in candidates:
            value = (candidate.value or "").strip()
            if not value:
                continue
            if candidate.confidence < min_confidence:
                continue
            if not self._value_matches_field_type(field_key, value):
                continue
            return {
                "value": value,
                "confidence": float(candidate.confidence),
                "source": str(candidate.memory_type or "rag"),
                "tier": int(candidate.metadata.get("tier", 99)),
            }
        return None

    @staticmethod
    def _value_matches_field_type(field_key: str, value: str) -> bool:
        key = f" {field_key} "
        val = value.strip()
        if not val:
            return False
        if " email " in key or key.strip() == "email":
            return bool(re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", val, re.IGNORECASE))
        if any(token in key for token in (" phone ", " sdt ", " dien_thoai ", " mobile ")):
            digits = re.sub(r"\D", "", val)
            return 8 <= len(digits) <= 15
        if " ngay " in key or " date " in key or " dob " in key:
            return bool(re.search(r"\d", val))
        return len(val) <= 800
