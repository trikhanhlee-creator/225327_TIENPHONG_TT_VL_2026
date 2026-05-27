"""
Convert legacy Word .doc files to .docx for python-docx and field detection.

Strategies (in order):
1. LibreOffice / OpenOffice headless (cross-platform)
2. Microsoft Word via COM on Windows (if pywin32 available)
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import tempfile
import time

from app.core.logger import logger


class DocConversionError(RuntimeError):
    """Raised when a .doc file cannot be converted to .docx."""


def _is_ole_doc(file_path: str) -> bool:
    """Best-effort check that file is a legacy OLE Word document."""
    try:
        import olefile

        if not olefile.isOleFile(file_path):
            return False
        with olefile.OleFileIO(file_path) as ole:
            return ole.exists("WordDocument")
    except Exception:
        return file_path.lower().endswith(".doc")


def _find_soffice() -> str | None:
    names = (
        "soffice",
        "libreoffice",
        "soffice.exe",
    )
    for name in names:
        found = shutil.which(name)
        if found:
            return found

    if sys.platform == "win32":
        program_roots = (
            os.environ.get("ProgramFiles", r"C:\Program Files"),
            os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)"),
        )
        for root in program_roots:
            candidate = os.path.join(root, "LibreOffice", "program", "soffice.exe")
            if os.path.isfile(candidate):
                return candidate
    return None


def _convert_with_libreoffice(doc_path: str, out_dir: str) -> str:
    soffice = _find_soffice()
    if not soffice:
        raise DocConversionError("LibreOffice chưa được cài đặt hoặc không có trong PATH.")

    abs_doc = os.path.abspath(doc_path)
    abs_out = os.path.abspath(out_dir)
    os.makedirs(abs_out, exist_ok=True)

    cmd = [
        soffice,
        "--headless",
        "--norestore",
        "--nologo",
        "--convert-to",
        "docx",
        "--outdir",
        abs_out,
        abs_doc,
    ]
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        raise DocConversionError("Chuyển đổi .doc quá thời gian chờ (120s).") from exc

    if result.returncode != 0:
        stderr = (result.stderr or result.stdout or "").strip()
        raise DocConversionError(stderr or "LibreOffice không chuyển đổi được file .doc.")

    base_name = os.path.splitext(os.path.basename(abs_doc))[0]
    expected = os.path.join(abs_out, f"{base_name}.docx")
    if os.path.isfile(expected):
        return expected

    raise DocConversionError("LibreOffice chạy xong nhưng không tạo file .docx.")


def _set_word_com_quiet(word) -> None:
    """Best-effort: hide Word UI; some installs disallow changing Visible."""
    for prop, value in (("Visible", False), ("DisplayAlerts", 0)):
        try:
            setattr(word, prop, value)
        except Exception:
            pass


def _get_word_application():
    """Attach to running Word or start a new instance (Windows only)."""
    import win32com.client  # type: ignore

    try:
        return win32com.client.GetObject(Class="Word.Application"), False
    except Exception:
        pass

    last_exc: Exception | None = None
    for attempt in range(3):
        try:
            word = win32com.client.gencache.EnsureDispatch("Word.Application")
            return word, True
        except Exception as exc:
            last_exc = exc
            if attempt < 2:
                time.sleep(1.5)
    raise DocConversionError(f"Không khởi động được Microsoft Word: {last_exc}")


def _convert_with_word_com(doc_path: str, out_path: str) -> str:
    if sys.platform != "win32":
        raise DocConversionError("Microsoft Word COM chỉ khả dụng trên Windows.")

    try:
        import win32com.client  # type: ignore  # noqa: F401
    except ImportError as exc:
        raise DocConversionError("Thiếu pywin32 để dùng Microsoft Word chuyển đổi .doc.") from exc

    abs_doc = os.path.abspath(doc_path)
    abs_out = os.path.abspath(out_path)
    os.makedirs(os.path.dirname(abs_out) or ".", exist_ok=True)

    word = None
    doc = None
    started_word = False
    try:
        word, started_word = _get_word_application()
        _set_word_com_quiet(word)
        doc = word.Documents.Open(abs_doc, ReadOnly=True)
        # 16 = wdFormatXMLDocument (.docx)
        doc.SaveAs2(abs_out, FileFormat=16)
        doc.Close(False)
        doc = None
    except Exception as exc:
        raise DocConversionError(f"Microsoft Word không chuyển đổi được file: {exc}") from exc
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if started_word and word is not None:
            try:
                word.Quit()
            except Exception:
                pass

    if not os.path.isfile(abs_out):
        raise DocConversionError("Word không tạo được file .docx đầu ra.")
    return abs_out


def _extract_with_antiword(doc_path: str) -> str | None:
    antiword = shutil.which("antiword")
    if not antiword:
        return None
    try:
        result = subprocess.run(
            [antiword, "-w", "0", os.path.abspath(doc_path)],
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )
    except (subprocess.TimeoutExpired, OSError):
        return None
    if result.returncode != 0:
        return None
    text = (result.stdout or "").strip()
    return text or None


def _extract_doc_plaintext_heuristic(doc_path: str) -> str:
    """Extract readable text from legacy .doc without Word/LibreOffice."""
    data: bytes
    try:
        import olefile

        with olefile.OleFileIO(doc_path) as ole:
            if ole.exists("WordDocument"):
                data = ole.openstream("WordDocument").read()
            else:
                with open(doc_path, "rb") as handle:
                    data = handle.read()
    except Exception:
        with open(doc_path, "rb") as handle:
            data = handle.read()

    chunks: list[str] = []
    i = 0
    length = len(data)
    while i < length - 1:
        if data[i] >= 0x20 and data[i] < 0x7F and data[i + 1] == 0:
            chars: list[str] = []
            while i < length - 1 and data[i + 1] == 0 and 0x20 <= data[i] < 0x7F:
                ch = chr(data[i])
                chars.append("\n" if ch in "\r\n\x07\x0b" else ch)
                i += 2
            if len(chars) >= 4:
                chunks.append("".join(chars))
            continue
        i += 1

    lines: list[str] = []
    seen: set[str] = set()
    for chunk in chunks:
        for raw_line in re.split(r"[\r\n]+", chunk):
            line = re.sub(r"\s+", " ", raw_line).strip()
            if len(line) < 2 or line in seen:
                continue
            seen.add(line)
            lines.append(line)
    return "\n".join(lines)


def extract_doc_plaintext(doc_path: str) -> str:
    """Best-effort plain text from .doc (antiword, then binary heuristic)."""
    text = _extract_with_antiword(doc_path)
    if text:
        return text
    return _extract_doc_plaintext_heuristic(doc_path)


def _create_docx_from_plaintext(text: str, out_path: str) -> str:
    from docx import Document

    doc = Document()
    for line in (text or "").splitlines():
        cleaned = line.strip()
        if cleaned:
            doc.add_paragraph(cleaned)
    if not doc.paragraphs:
        doc.add_paragraph("")
    abs_out = os.path.abspath(out_path)
    os.makedirs(os.path.dirname(abs_out) or ".", exist_ok=True)
    doc.save(abs_out)
    return abs_out


def _convert_with_plaintext_fallback(doc_path: str, out_path: str) -> str:
    text = extract_doc_plaintext(doc_path)
    if len(text.strip()) < 8:
        raise DocConversionError(
            "Không trích xuất được nội dung từ file .doc. "
            "Hãy cài LibreOffice hoặc Microsoft Word, hoặc lưu file dạng .docx."
        )
    logger.warning(
        "Using plaintext fallback for .doc conversion (layout may differ): %s",
        os.path.basename(doc_path),
    )
    return _create_docx_from_plaintext(text, out_path)


def convert_doc_to_docx(doc_path: str, output_path: str | None = None) -> str:
    """
    Convert .doc to .docx.

    Returns absolute path to the generated .docx file.
  """
    if not doc_path or not os.path.isfile(doc_path):
        raise DocConversionError(f"File không tồn tại: {doc_path}")

    ext = os.path.splitext(doc_path)[1].lower()
    if ext != ".doc":
        raise DocConversionError(f"Không phải file .doc: {ext}")

    if not _is_ole_doc(doc_path):
        raise DocConversionError("File không phải định dạng Word .doc hợp lệ.")

    if output_path:
        out_path = os.path.abspath(output_path)
        out_dir = os.path.dirname(out_path) or os.path.dirname(os.path.abspath(doc_path))
    else:
        out_dir = os.path.dirname(os.path.abspath(doc_path))
        out_path = os.path.join(out_dir, f"{os.path.splitext(os.path.basename(doc_path))[0]}.docx")

    errors: list[str] = []

    try:
        converted = _convert_with_libreoffice(doc_path, out_dir)
        if output_path and os.path.abspath(converted) != out_path:
            shutil.move(converted, out_path)
            return out_path
        return converted
    except DocConversionError as exc:
        errors.append(f"LibreOffice: {exc}")

    if sys.platform == "win32":
        try:
            return _convert_with_word_com(doc_path, out_path)
        except DocConversionError as exc:
            errors.append(str(exc))

    try:
        return _convert_with_plaintext_fallback(doc_path, out_path)
    except DocConversionError as exc:
        errors.append(str(exc))

    hint = (
        "Cài LibreOffice (https://www.libreoffice.org/) hoặc mở file trong Word và lưu lại dạng .docx."
    )
    raise DocConversionError(
        "Không thể đọc file .doc. " + " | ".join(errors) + f" {hint}"
    )


def ensure_docx_for_processing(file_path: str) -> str:
    """Return a .docx path suitable for python-docx (convert .doc if needed)."""
    ext = os.path.splitext(file_path or "")[1].lower()
    if ext == ".docx":
        return file_path
    if ext == ".doc":
        return convert_doc_to_docx(file_path)
    return file_path
