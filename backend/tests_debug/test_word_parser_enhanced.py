#!/usr/bin/env python3
"""Regression checks for enhanced Word parser (tables, checkboxes, multi-row)."""

import os
import sys
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document

from app.services.file_parser import FileParserFactory


def _make_temp_docx(build_fn):
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc = Document()
        build_fn(doc)
        doc.save(path)
        return path
    except Exception:
        if os.path.exists(path):
            os.remove(path)
        raise


def _parse_fields(path: str):
    parser = FileParserFactory.create_parser(path)
    return parser.parse()


def test_choice_field_from_checkboxes() -> None:
    def build(doc: Document) -> None:
        doc.add_paragraph("Giới tính: [ ] Nam   [ ] Nữ   [ ] Khác")

    path = _make_temp_docx(build)
    try:
        fields = _parse_fields(path)
    finally:
        os.remove(path)

    choice_fields = [f for f in fields if f.field_type == "choice"]
    assert choice_fields, f"Expected choice field, got={[f.to_dict() for f in fields]}"
    field = choice_fields[0]
    assert field.label.lower().startswith("giới tính") or "gioi" in field.name
    assert set(field.options) >= {"Nam", "Nữ"}


def test_table_with_header_and_multi_row() -> None:
    def build(doc: Document) -> None:
        table = doc.add_table(rows=4, cols=3)
        headers = ["STT", "Trường", "Năm tốt nghiệp"]
        for idx, title in enumerate(headers):
            table.cell(0, idx).text = title
        rows = [
            ("1", "ĐH Bách Khoa", "................"),
            ("2", "THPT Nguyễn Huệ", "................"),
            ("3", "THCS Lê Lợi", "................"),
        ]
        for row_idx, values in enumerate(rows, start=1):
            for col_idx, value in enumerate(values):
                table.cell(row_idx, col_idx).text = value

    path = _make_temp_docx(build)
    try:
        fields = _parse_fields(path)
    finally:
        os.remove(path)

    names = [f.name for f in fields]
    assert any("_1" in name for name in names), f"Expected row suffix fields, got={names}"
    assert any("_2" in name for name in names), f"Expected row suffix fields, got={names}"
    assert len(fields) >= 3


def test_document_order_paragraph_before_table() -> None:
    def build(doc: Document) -> None:
        doc.add_paragraph("Họ và tên: ..............................")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Email"
        table.cell(0, 1).text = "......................"

    path = _make_temp_docx(build)
    try:
        fields = _parse_fields(path)
    finally:
        os.remove(path)

    labels = [f.label for f in fields]
    assert "Họ và tên" in labels
    assert "Email" in labels
    ho_idx = next(i for i, f in enumerate(fields) if "Họ" in f.label or "ten" in f.name)
    email_idx = next(i for i, f in enumerate(fields) if f.label == "Email")
    assert ho_idx < email_idx, f"Paragraph field should appear before table field: {labels}"


if __name__ == "__main__":
    test_choice_field_from_checkboxes()
    test_table_with_header_and_multi_row()
    test_document_order_paragraph_before_table()
    print("All enhanced Word parser tests passed.")
